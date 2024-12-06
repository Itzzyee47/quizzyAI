import pyrebase
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document
import requests
import json
import re, os
from dotenv import load_dotenv

load_dotenv()

key = os.environ["API_KEY"]
 
firebaseConfig ={
  "apiKey": key,
  "authDomain": "school-management-e7a71.firebaseapp.com",
  "databaseURL": "https://school-management-e7a71-default-rtdb.firebaseio.com",
  "projectId": "school-management-e7a71",
  "storageBucket": "school-management-e7a71.appspot.com",
  "messagingSenderId": "1065954376765",
  "appId": "1:1065954376765:web:b42912ed618c5a7c11be54",
  "measurementId": "G-ZMLC7LKVKC"
}

firebase_storage = pyrebase.initialize_app(firebaseConfig)
storage = firebase_storage.storage()

# Get the blob (file) from Firebase Storage
test_file = "notes/1732870719339_DIGITAL ELECTRONICS.docx"  #"https://firebasestorage.googleapis.com/v0/b/lsachatbot.appspot.com/o/notes%2FCourseOutlineForBlender.docx?alt=media&token=15c78df0-157b-4bbf-8b1e-dc3d85dc34fe"
test2_file = "notes/neralNetworks.docx" 
test3_file = "notes/Newwhitepaper_Foundational Large Language models & text generation.pdf"

def extract_file_type(url):
    """extract file type."""
    # Determine file type by extension
    if url.endswith(".pdf"):
        file_type = "PDF"
        return file_type
    elif url.endswith(".docx"):
        file_type = "DOCX"
        return file_type
    else:
        raise ValueError("Unsupported file type. Only PDF and DOCX are supported.")

# Extract text from pdf files..........
def extract_text_from_pdf(file_stream):
    """Extract text from a PDF file."""
    reader = PdfReader(file_stream)
    all_text = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            all_text.append(text)
    return "\n".join(all_text)

# Extract text from docx files..........
def extract_text_from_docx(file_stream):
    """Extract text from a DOCX file."""
    doc = Document(file_stream)
    all_text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Skip empty paragraphs
            all_text.append(paragraph.text.strip())
    
    all_text = [p.text for p in doc.paragraphs if p.text.strip()]
    # Combine all text into a single string (optional)
    return "\n".join(all_text)

def get_course_notes(file_path):
    # Download the file from Firebase Storage as bytes
    print('getting document')
    # Get file type....
    file_type = extract_file_type(file_path)
   # Download file as bytes
    try:
        # Get the public URL of the file
        file_url = storage.child(file_path).get_url(None)
        print(f"Downloading file from URL: {file_url}")

        # Download the file content using requests
        response = requests.get(url=file_url)
        #print(f"HTTP Status Code: {response.status_code}")
        response.raise_for_status()  # Raise an exception for HTTP errors
        
        # Load the content into a BytesIO stream
        document_stream = BytesIO(response.content)
        # Step 3: Load the word docx document using python-docx
        print('Reading document...')
        if file_type == "DOCX":
            full_text = extract_text_from_docx(document_stream)
        elif file_type == "PDF":
            full_text = extract_text_from_pdf(document_stream)
        
        return full_text

    except Exception as e:
        return f"Error fetching or processing document: {e}"

#get_course_notes("notes/Documentation of cost estimate of using the gemini api per year.docx")

def extract_quiz_json(data):
    try:
            # Use a regular expression to extract the JSON string between triple backticks
            json_match = re.search(r"```json\s*(.*?)\s*```", data, re.DOTALL)
            if json_match:
                # Parse the JSON string into a Python dictionary
                quiz_data = json.loads(json_match.group(1))
                return quiz_data
            else:
                print("No JSON data found in the response.")
    except Exception as e:
        print(f"Failed to parse JSON data: {e}")
        
def generate_quiz_with_answers(notes):
    json_example = str( {    "quiz": [        {           "question": "What is AI?",           "options": {"Option 1", "Option 2", "Option 3", "Option 4"},  "answer": "Option 1"       },       {      "question": "Define ML.",            "options": {"Option 1", "Option 2", "Option 3", "Option 4"},           "answer": "Option 1"        }    ]} )
    
    payload = dict(message=f"Generate me a 10-question multiple-choice quiz based on the sumarized module's content of the most relevant contexts in the module. Each question must have 4 options, with only one correct answer. Heres the module notes {notes}, return only the questions and options in a json format i can easily read with code and must be identical to this JSON object in the following structure: {json_example}")
    r = requests.post('https://genzylla.onrender.com/getResponds', data=payload)
    print(f"The responds data type is: {type(r.text)}")
    res = r.json()
    quiz_result = extract_quiz_json(res['answer'])
    return quiz_result